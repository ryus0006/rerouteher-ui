import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                      f'<w:top w:w="{top}" w:type="dxa"/>'
                      f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                      f'<w:left w:w="{left}" w:type="dxa"/>'
                      f'<w:right w:w="{right}" w:type="dxa"/>'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

def set_cell_left_border(cell, hex_color="DE8BA8", size="28"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="{hex_color}"/>'
                        f'<w:top w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'<w:bottom w:val="none"/>'
                        f'</w:tcBorders>')
    tcPr.append(borders)

def set_table_borders(table, hex_color="D6D2E0"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{hex_color}"/>'
                        f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="3B4B7C"/>'
                        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
                        f'<w:insideV w:val="none"/>'
                        f'<w:left w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'</w:tblBorders>')
    tblPr.append(borders)

def add_screenshot_figure(doc, img_path, caption_text, width_in=6.2):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(3)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width_in))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run(caption_text)
        r_cap.font.name = 'Plus Jakarta Sans'
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x5A, 0x60, 0x86)

def build_word_document():
    doc = Document()

    # Page Margins (Normal 1 inch / 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Plus Jakarta Sans'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x26, 0x2B, 0x4A) # #262B4A
    normal_style.paragraph_format.line_spacing = 1.25
    normal_style.paragraph_format.space_after = Pt(4)

    # Palette
    C_NAVY = RGBColor(0x26, 0x2B, 0x4A)      # #262B4A
    C_BLUE = RGBColor(0x3B, 0x4B, 0x7C)      # #3B4B7C
    C_PINK = RGBColor(0xDE, 0x8B, 0xA8)      # #DE8BA8
    C_MINT = RGBColor(0x23, 0x5C, 0x41)      # #235C41
    C_AMBER = RGBColor(0x96, 0x54, 0x0D)     # #96540D
    C_MUTED = RGBColor(0x5A, 0x60, 0x86)     # #5A6086

    # ==================== HEADER & TITLE ====================
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("📱 09 · Low-Fidelity & High-Level (Hi-Fi) Interactive Prototypes")
    run_title.font.name = 'Plus Jakarta Sans'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = C_NAVY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("ReRouteHer — AI-Powered Skill Readiness & Career Re-entry Platform")
    run_sub.font.name = 'Plus Jakarta Sans'
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = C_PINK

    # Metadata Callout Card
    meta_table = doc.add_table(rows=1, cols=1)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(6.5)
    
    cell = meta_table.cell(0, 0)
    set_cell_background(cell, "FBF0F4")
    set_cell_left_border(cell, "DE8BA8", "32")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    mp = cell.paragraphs[0]
    mp.paragraph_format.space_after = Pt(2)
    
    r = mp.add_run("Course / Team Code: ")
    r.bold = True
    r.font.color.rgb = C_NAVY
    mp.add_run("5120-TM07\n")
    
    r = mp.add_run("Figma Canvas: ")
    r.bold = True
    r.font.color.rgb = C_NAVY
    mp.add_run("https://www.figma.com/design/wwl5kUAg2RJF8anRURdf1n/5120-TM07\n")
    
    r = mp.add_run("Live Interactive Prototype: ")
    r.bold = True
    r.font.color.rgb = C_NAVY
    mp.add_run("https://prototype.curl.my/\n")
    
    r = mp.add_run("GitHub Repositories: ")
    r.bold = True
    r.font.color.rgb = C_NAVY
    mp.add_run("ryus0006/rerouteher-ui & ncm233/reroutehers-prototype")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ==================== SECTION 1: EXECUTIVE SUMMARY ====================
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("1. Executive Summary & Prototyping Objectives")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.add_run(
        "Returning to formal employment after a prolonged caregiving career break is fraught with severe emotional and cognitive hurdles: "
        "returners frequently experience "
    )
    r = p.add_run("imposter syndrome")
    r.bold = True
    p.add_run(
        ", struggle to articulate their transferable capabilities, and encounter a "
    )
    r = p.add_run('"wall of 20+ overwhelming job requirements"')
    r.bold = True
    p.add_run(" on traditional job boards.")

    p2 = doc.add_paragraph()
    r = p2.add_run("ReRouteHer")
    r.bold = True
    r.font.color.rgb = C_PINK
    p2.add_run(
        " addresses this gap through an empathy-led, data-backed career re-entry readiness engine. "
        "The prototyping phase followed an iterative double-diamond design lifecycle, progressing from conceptual low-fidelity (Lo-Fi) structural wireframes "
        "to a high-level, interactive high-fidelity (Hi-Fi) prototype."
    )

    # ==================== SECTION 2: LOW-FIDELITY WIREFRAMES ====================
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("2. Low-Fidelity (Lo-Fi) Prototypes & Wireframe Architecture")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("2.1 Design Rationale & Information Hierarchy")
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    p = doc.add_paragraph()
    p.add_run("The initial low-fidelity wireframes established the foundational four-step sequential journey, deliberately stripping away visual embellishments to validate three core design hypotheses:")
    
    bullets = [
        ("Zero-Friction Intake: ", "Eliminating multi-page tedious questionnaires in favor of an automated 2-step input (CV document parsing + natural language career break description)."),
        ("Reflective vs. Exploratory Separation: ", "Separating the historical baseline (Skill Snapshot — where she is coming from) from the aspirational target (Target Role & Gap — where she wants to aim)."),
        ("Anti-Overwhelm Focus: ", "Guaranteeing a strict maximum cap of 3 actionable focus areas, replacing demoralizing comprehensive requirement checklists.")
    ]
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(3)
        r = bp.add_run(b_title)
        r.bold = True
        r.font.color.rgb = C_NAVY
        bp.add_run(b_desc)

    # Table: Lo-Fi Breakdown
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("2.2 Screen-by-Screen Lo-Fi Breakdown")
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    lofi_table_data = [
        ("Stage", "Lo-Fi Screen Name", "Key Functional Modules & Layout Strategy", "User Experience Intent"),
        ("01", "Landing Page", "• Minimalist top navigation with brand mark\n• Hero header with bold value proposition\n• Prominent primary CTA button ('Start Your Journey')\n• 3-column value cards (Break Experience, Weighted Score, Top 3 Focus)", "Instills immediate confidence and sets clear expectations of zero manual friction."),
        ("02", "Profile & Experience Intake", "• Drag-and-drop CV upload dropzone (PDF/DOCX)\n• Duration selector (0.5 to 15 years)\n• Free-text multi-line textarea with example suggestion chips", "Bypasses repetitive form fields, allowing returners to express real-world experiences in their own words."),
        ("03", "Baseline Skill Snapshot", "• Non-locking 'Background Baseline' header\n• Two-column skill container: Extracted CV Skills vs. Reframed Break Skills\n• O*NET crosswalk indicator badge", "Validates both professional achievements and caregiving competencies on an equal footing."),
        ("04", "Target Role & Gap Analysis", "• Interactive target role selector tabs\n• Semicircular readiness gauge displaying current baseline %\n• Capped Top 3 focus areas with projected readiness boost tags\n• 'Explore Upskilling Paths' CTA", "Replaces failure anxiety with an actionable, bite-sized growth plan.")
    ]

    t = doc.add_table(rows=len(lofi_table_data), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t)

    col_widths = [Inches(0.6), Inches(1.5), Inches(2.6), Inches(1.8)]

    for row_idx, row_data in enumerate(lofi_table_data):
        row = t.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            c = row.cells[col_idx]
            c.width = col_widths[col_idx]
            set_cell_margins(c, top=90, bottom=90, left=110, right=110)
            cp = c.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            
            if row_idx == 0:
                set_cell_background(c, "262B4A")
                r = cp.add_run(text)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9.5)
            else:
                if row_idx % 2 == 1:
                    set_cell_background(c, "FDFBFC")
                else:
                    set_cell_background(c, "F8F5FA")
                r = cp.add_run(text)
                r.font.size = Pt(9)
                r.font.color.rgb = C_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== SECTION 3: HIGH-FIDELITY PROTOTYPE (DIRECT SCREENSHOTS) ====================
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("3. High-Fidelity (Hi-Fi) Interactive Prototype (Direct Prototype Captures)")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.add_run("The high-fidelity prototype transforms the validated wireframes into a soothing, empowering, state-of-the-art interactive web application. Below are direct, unedited high-resolution captures of the live prototype implementation.")

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("3.1 Design System & Visual Tokens")
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    # Design Tokens Table
    color_table_data = [
        ("Token Name", "Hex Code", "Tailwind / CSS Var", "Semantics & Psychological Function"),
        ("Primary Blush Pink", "#DE8BA8", "--pink-500", "Metamorphosis, empathy, primary gradient CTA button"),
        ("Soft Lavender", "#B4A2D4", "--violet-400", "Transmutation, cognitive calming, ethereal depth"),
        ("Periwinkle Blue", "#7E92CA", "--blue-600", "Trust, professional stability, technical skill tags"),
        ("Mint Green", "#337857", "--mint-600", "O*NET validated break achievements, high confidence badge"),
        ("Amber Gold", "#96540D", "--amber-700", "Priority focus area uplift badges (+% gain)"),
        ("Midnight Ink", "#262B4A", "--ink", "High-contrast primary typography & UI structure"),
        ("Frosted Canvas", "#FCF8FA", "--grad-soft", "Multi-stop ethereal soft gradient page background")
    ]

    t_color = doc.add_table(rows=len(color_table_data), cols=4)
    t_color.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_color.autofit = False
    set_table_borders(t_color)
    c_widths = [Inches(1.5), Inches(1.0), Inches(1.3), Inches(2.7)]

    for row_idx, row_data in enumerate(color_table_data):
        row = t_color.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            c = row.cells[col_idx]
            c.width = c_widths[col_idx]
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            cp = c.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_background(c, "262B4A")
                r = cp.add_run(text)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
            else:
                if row_idx % 2 == 1:
                    set_cell_background(c, "FDFBFC")
                else:
                    set_cell_background(c, "F8F5FA")
                r = cp.add_run(text)
                r.font.size = Pt(8.5)
                r.font.color.rgb = C_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Direct Screen Captures & Specifications
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    r = h2.add_run("3.2 Direct Screen Captures & Interactive Specifications")
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    # 1. Landing Page
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(2)
    r = h3.add_run("Screen 1 · Landing Page (E1)")
    r.bold = True
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)

    add_screenshot_figure(doc, "screenshots/01_landing_page.png", "Figure 1: Direct Capture — ReRouteHer Landing Page (E1) with Hero Butterfly Artwork & 3-Step Journey Rail")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Hero Artwork: ")
    r.bold = True
    p.add_run("Ethereal multi-layered butterfly oil painting with radial alpha-mask gradient blending (85% opacity).")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Journey Rail: ")
    r.bold = True
    p.add_run("3-step interactive visual progress stepper (Upload CV ➔ Describe Break ➔ See Fit & Top Gaps).")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Parallax Dynamics: ")
    r.bold = True
    p.add_run("Smooth GSAP orbital physics and soft floating star highlights.")

    # 2a. Step 1: Upload CV
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(2)
    r = h3.add_run("Screen 2a · Step 1: Upload CV (E2a)")
    r.bold = True
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)

    add_screenshot_figure(doc, "screenshots/02_cv_upload.png", "Figure 2: Direct Capture — Step 1: Upload CV (E2a) Drag-and-Drop Dropzone & Sample Resumes")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("File Dropzone: ")
    r.bold = True
    p.add_run("Drag-and-drop file upload supporting .pdf and .docx (up to 10MB) with instant client-side verification.")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("1-Click Sample CV Loaders: ")
    r.bold = True
    p.add_run("Pre-configured analyst and designer sample resumes for immediate friction-free evaluation.")

    # 2b. Step 2: Career Break Intake
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(2)
    r = h3.add_run("Screen 2b · Step 2: Career Break Intake (E2b)")
    r.bold = True
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)

    add_screenshot_figure(doc, "screenshots/03_career_break.png", "Figure 3: Direct Capture — Step 2: Career Break Intake (E2b) Duration Slider & Natural Language Textarea")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Duration Slider: ")
    r.bold = True
    p.add_run("Intuitive range slider (0.5 to 15 years) with dynamic mint tag preview.")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Natural Language Textarea: ")
    r.bold = True
    p.add_run("Free-text input capturing caregiving, budgeting, volunteering, and self-study, accompanied by 4 one-tap example tags (+Childcare, +Budgeting, +Volunteering, +Self-study).")

    # 3. Step 3: Skill Snapshot Baseline
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(2)
    r = h3.add_run("Screen 3 · Step 3: Skill Snapshot Baseline (E3)")
    r.bold = True
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)

    add_screenshot_figure(doc, "screenshots/04_skill_snapshot.png", "Figure 4: Direct Capture — Step 3: Skill Snapshot Baseline (E3) Occupation Baseline & Two-Column Skill Inventory")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Occupation Baseline Line: ")
    r.bold = True
    p.add_run("Features an explicit headline based on the backend reranker: 'Based on your story, you're closest to Operation Research Analyst' with a 'High confidence match' mint badge.")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("From your CV: ")
    r.bold = True
    p.add_run("Extracted career competencies rendered as compact pill chips ('SkillChip') with hover evidence tooltips and a 'Show all / Show fewer' collapse toggle.")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("From your career break: ")
    r.bold = True
    p.add_run("O*NET-reframed domestic and community skills (Active Listening, Social Perceptiveness, Time Management, Coordination, Management of Financial Resources).")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("O*NET Crosswalk Bridge: ")
    r.bold = True
    p.add_run("Informational banner illustrating automated NLP translation from everyday tasks to US Dept. of Labor taxonomies.")

    # 4. Step 4: Target Role & Gap Analysis
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(2)
    r = h3.add_run("Screen 4 · Step 4: Target Role & Gap Analysis (E4)")
    r.bold = True
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)

    add_screenshot_figure(doc, "screenshots/05_target_role_gap.png", "Figure 5: Direct Capture — Step 4: Target Role & Gap Analysis (E4) 62.6% Readiness Gauge & Top 3 Priority Focus Areas")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Interactive Role Selector Pills: ")
    r.bold = True
    p.add_run("Operation Research Analyst (Closest match — active), Data Analyst, Management Information Systems (MIS) Analyst.")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("210° Arc Readiness Gauge: ")
    r.bold = True
    p.add_run("Dynamic SVG sweep gauge indicating '62.6% READY TODAY' paired with a projected readiness card: '62.6% today → 84.3% after your focus areas'.")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Missing for this role (Capped Top 3 Focus Areas): ")
    r.bold = True
    p.add_run("1. Mathematics (O*NET Skill) [+6.7%], 2. Use AI Assistants for Everyday Work Tasks [+7.5%], 3. Check and Verify AI Output [+7.5%].")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Importance-Weighted Formula Card: ")
    r.bold = True
    p.add_run("Transparently explains the scoring mechanism to demystify readiness percentages.")

    # 4b. Role Variant: Data Analyst
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(2)
    r = h3.add_run("Screen 4 (Variant) · Target Role Switching (Data Analyst)")
    r.bold = True
    r.font.color.rgb = C_BLUE
    r.font.size = Pt(11)

    add_screenshot_figure(doc, "screenshots/06_role_data_analyst.png", "Figure 6: Direct Capture — Target Role Variant Switching to Data Analyst (71.4% Baseline ➔ 88.9% Target)")

    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run("Dynamic Live Recalculation: ")
    r.bold = True
    p.add_run("Demonstrates instant recalculation when switching to Data Analyst (71.4% Baseline ➔ 88.9% Target) with tailored focus areas (SQL Optimization, AI Assistants, Tableau/Power BI).")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==================== SECTION 4: USABILITY & ITERATION LOG ====================
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("4. Usability Testing & Iterative Design Refinement Log")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    iter_table_data = [
        ("Iteration Phase", "Identified User Pain Point / Feedback", "Implemented Design Solution in Hi-Fi Prototype"),
        ("Lo-Fi Concept", "Generic role selection felt restrictive and forced returners into boxes prematurely.", "Introduced the Read-Only Skill Snapshot step as a reflective baseline before prompting role selection."),
        ("Initial Hi-Fi (v1.0)", "AI literacy skills flooded the focus list, crowding out essential domain technical skills.", "Engineered the pickFocusAreas algorithm: guarantees 1 AI-literacy slot and reserves remaining slots for core domain role skills."),
        ("User Feedback (v1.1)", "Long CV skill lists created vertical scrolling clutter and cognitive fatigue.", "Created compact pill chips ('SkillChip') with hover evidence tooltips and a 'Show all / Show fewer' collapse toggle."),
        ("Team Review (v2.0)", "Perceived disconnect between raw skill count (e.g. 7 of 10) and percentage readiness (62.6%).", "Added the Importance-Weighted Formula Card, clearly explaining weighting factors and O*NET skill benchmarks.")
    ]

    t_iter = doc.add_table(rows=len(iter_table_data), cols=3)
    t_iter.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_iter.autofit = False
    set_table_borders(t_iter)
    it_widths = [Inches(1.4), Inches(2.5), Inches(2.6)]

    for row_idx, row_data in enumerate(iter_table_data):
        row = t_iter.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            c = row.cells[col_idx]
            c.width = it_widths[col_idx]
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            cp = c.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_background(c, "262B4A")
                r = cp.add_run(text)
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
            else:
                if row_idx % 2 == 1:
                    set_cell_background(c, "FDFBFC")
                else:
                    set_cell_background(c, "F8F5FA")
                r = cp.add_run(text)
                r.font.size = Pt(8.5)
                r.font.color.rgb = C_NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==================== SECTION 5: ACCESS & VERIFICATION ====================
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("5. Deliverable Links & Interactive Prototype Access")
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    final_box = doc.add_table(rows=1, cols=1)
    final_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    final_box.autofit = False
    final_box.columns[0].width = Inches(6.5)
    
    cell = final_box.cell(0, 0)
    set_cell_background(cell, "DEF3E7")
    set_cell_left_border(cell, "235C41", "32")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    fp = cell.paragraphs[0]
    fp.paragraph_format.space_after = Pt(4)
    
    r = fp.add_run("🌐 Live Deployed Prototype URL: ")
    r.bold = True
    r.font.color.rgb = C_MINT
    fp.add_run("https://prototype.curl.my/\n\n")

    r = fp.add_run("🎨 Figma Interactive Design Canvas: ")
    r.bold = True
    r.font.color.rgb = C_MINT
    fp.add_run("https://www.figma.com/design/wwl5kUAg2RJF8anRURdf1n/5120-TM07?node-id=0-1&t=8vvvEvjkLIAUzMhL-1\n\n")

    r = fp.add_run("💻 Local Development & Execution Commands:\n")
    r.bold = True
    r.font.color.rgb = C_MINT
    fp.add_run("• Install dependencies: npm install\n• Start development server: npm run dev (http://localhost:5173)\n• Run unit test suite: npm run test:unit (26 passing tests)")

    # Save documents
    out_path1 = "09_lofi_and_hifi_prototypes.docx"
    out_path2 = "09_Lo-Fi_and_Hi-Fi_Prototypes_5120-TM07.docx"
    doc.save(out_path1)
    doc.save(out_path2)
    print(f"Successfully generated {out_path1} and {out_path2}")

if __name__ == '__main__':
    build_word_document()
